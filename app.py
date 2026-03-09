"""
LangChain Summarization + RAG Chat App
=======================================
Features:
  - Summarize YouTube videos and web pages (map-reduce)
  - Multi-URL loading: add several URLs and query across all of them
  - URL history: previously loaded URLs are cached; switch without re-processing
  - Conversation memory: the chat remembers prior turns in the session
  - Download summary as TXT / PDF / DOCX
  - RAG with FAISS + HuggingFace embeddings

Dependencies:
    pip install streamlit langchain-core langchain-groq langchain-community
                langchain-text-splitters langchain-huggingface python-dotenv
                validators youtube-transcript-api unstructured requests
                fpdf2 python-docx faiss-cpu sentence-transformers
"""

import io
import logging
import re
from typing import Optional

import validators
import streamlit as st
from dotenv import load_dotenv
from fpdf import FPDF
from docx import Document as DocxDocument

from langchain_core.prompts import PromptTemplate
from langchain_core.documents import Document
from langchain_groq import ChatGroq
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import UnstructuredURLLoader
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import (
    TranscriptsDisabled,
    NoTranscriptFound,
    VideoUnavailable,
)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

DEFAULT_MODEL = "llama-3.1-8b-instant"
EMBED_MODEL   = "sentence-transformers/all-MiniLM-L6-v2"
CHUNK_SIZE    = 2000
CHUNK_OVERLAP = 200
MAX_TOKENS    = 1024
RAG_TOP_K     = 4   # retrieved chunks per question
MEMORY_TURNS  = 6   # how many past conversation turns to include in context

# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------

MAP_PROMPT = PromptTemplate(
    input_variables=["text"],
    template=(
        "Write a concise summary of the following text, preserving key facts and insights.\n\n"
        "TEXT:\n{text}\n\n"
        "CONCISE SUMMARY:"
    ),
)

COMBINE_PROMPT = PromptTemplate(
    input_variables=["text"],
    template=(
        "You are given partial summaries from a longer document.\n"
        "Combine them into one coherent, well-structured final summary (~300 words).\n"
        "Do not repeat yourself. Focus on the most important points.\n\n"
        "PARTIAL SUMMARIES:\n{text}\n\n"
        "FINAL SUMMARY:"
    ),
)

# Includes conversation history so the model can reference prior turns
RAG_PROMPT = PromptTemplate(
    input_variables=["context", "history", "question"],
    template=(
        "You are a helpful assistant. Answer using ONLY the context provided.\n"
        "If the answer is not in the context, say: 'I could not find that in the loaded content.'\n"
        "Be concise and direct. You may reference prior conversation turns if relevant.\n\n"
        "CONTEXT (retrieved from loaded URLs):\n{context}\n\n"
        "CONVERSATION HISTORY:\n{history}\n\n"
        "CURRENT QUESTION:\n{question}\n\n"
        "ANSWER:"
    ),
)

# ---------------------------------------------------------------------------
# LLM & Embeddings
# ---------------------------------------------------------------------------

def get_llm(api_key: str) -> ChatGroq:
    """Initialise and return the Groq LLM."""
    return ChatGroq(
        model=DEFAULT_MODEL,
        api_key=api_key,
        temperature=0.3,
        max_tokens=MAX_TOKENS,
    )


@st.cache_resource(show_spinner="Loading embedding model…")
def get_embeddings() -> HuggingFaceEmbeddings:
    """Load the embedding model once per session and keep it cached."""
    return HuggingFaceEmbeddings(model_name=EMBED_MODEL)


# ---------------------------------------------------------------------------
# Document Loading
# ---------------------------------------------------------------------------

def _extract_video_id(url: str) -> Optional[str]:
    patterns = [
        r"(?:v=|\/)([0-9A-Za-z_-]{11})",
        r"(?:youtu\.be\/)([0-9A-Za-z_-]{11})",
        r"(?:embed\/)([0-9A-Za-z_-]{11})",
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


def _load_youtube(url: str) -> list[Document]:
    """Fetch a YouTube transcript via youtube-transcript-api (v1.x)."""
    video_id = _extract_video_id(url)
    if not video_id:
        raise ValueError(f"Could not extract a video ID from: {url}")

    logger.info("Fetching transcript for video ID: %s", video_id)
    try:
        api = YouTubeTranscriptApi()
        try:
            fetched = api.fetch(video_id, languages=["en", "en-US", "en-GB"])
        except NoTranscriptFound:
            logger.warning("No English transcript — trying any available language.")
            fetched = api.fetch(video_id)

        full_text = " ".join(s.text for s in fetched.snippets)

    except VideoUnavailable:
        raise ValueError("This video is unavailable or private.")
    except TranscriptsDisabled:
        raise ValueError("Transcripts are disabled for this video.")
    except NoTranscriptFound:
        raise ValueError("No captions found for this video in any language.")
    except Exception as exc:
        raise RuntimeError(f"Failed to fetch YouTube transcript: {exc}") from exc

    if not full_text.strip():
        raise ValueError("The transcript was empty — nothing to process.")

    return [Document(page_content=full_text, metadata={"source": url, "video_id": video_id})]


def load_documents(url: str) -> list[Document]:
    """Load content from a YouTube URL or a generic web page."""
    try:
        if "youtube.com" in url or "youtu.be" in url:
            return _load_youtube(url)

        logger.info("Loading web content: %s", url)
        loader = UnstructuredURLLoader(
            urls=[url],
            ssl_verify=False,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/120.0 Safari/537.36"
                )
            },
        )
        docs = loader.load()
        if not docs:
            raise ValueError("No content could be extracted from the provided URL.")
        return docs

    except (ValueError, RuntimeError):
        raise
    except Exception as exc:
        logger.exception("Failed to load documents from %s", url)
        raise RuntimeError(f"Could not load content: {exc}") from exc


def split_documents(docs: list[Document]) -> list[Document]:
    """Split documents into LLM-sized chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunks = splitter.split_documents(docs)
    logger.info("Split into %d chunk(s)", len(chunks))
    return chunks


# ---------------------------------------------------------------------------
# Summarization
# ---------------------------------------------------------------------------

def run_map_reduce(llm: ChatGroq, chunks: list[Document]) -> str:
    """Map-reduce summarization — no external chain dependency."""
    summaries: list[str] = []
    for i, chunk in enumerate(chunks):
        logger.info("Summarizing chunk %d / %d", i + 1, len(chunks))
        resp = llm.invoke(MAP_PROMPT.format(text=chunk.page_content))
        summaries.append(resp.content.strip())

    combined = "\n\n---\n\n".join(summaries)
    final = llm.invoke(COMBINE_PROMPT.format(text=combined))
    return final.content.strip()


# ---------------------------------------------------------------------------
# RAG
# ---------------------------------------------------------------------------

def build_vector_store(chunks: list[Document]) -> FAISS:
    """Embed chunks and return an in-memory FAISS vector store."""
    return FAISS.from_documents(chunks, get_embeddings())


def merge_vector_stores(stores: list[FAISS]) -> FAISS:
    """
    Merge multiple FAISS indexes into one so queries run across
    all loaded URLs simultaneously.
    """
    base = stores[0]
    for store in stores[1:]:
        base.merge_from(store)
    return base


def _build_history_string(chat_history: list[dict]) -> str:
    """
    Format the last MEMORY_TURNS turns of chat history into a plain
    text block the LLM can read as conversation context.
    """
    recent = chat_history[-(MEMORY_TURNS * 2):]   # user+assistant pairs
    lines = []
    for msg in recent:
        role = "User" if msg["role"] == "user" else "Assistant"
        lines.append(f"{role}: {msg['content']}")
    return "\n".join(lines) if lines else "No prior conversation."


def answer_question(
    llm: ChatGroq,
    vector_store: FAISS,
    question: str,
    chat_history: list[dict],
) -> tuple[str, list[Document]]:
    """
    Retrieve top-K relevant chunks and answer the question using both
    the retrieved context and prior conversation history.

    Returns:
        (answer_text, source_chunks)
    """
    retriever = vector_store.as_retriever(
        search_type="similarity",
        search_kwargs={"k": RAG_TOP_K},
    )
    relevant_chunks = retriever.invoke(question)
    context  = "\n\n---\n\n".join(c.page_content for c in relevant_chunks)
    history  = _build_history_string(chat_history)
    response = llm.invoke(RAG_PROMPT.format(
        context=context,
        history=history,
        question=question,
    ))
    return response.content.strip(), relevant_chunks


# ---------------------------------------------------------------------------
# Session-state helpers
# ---------------------------------------------------------------------------

def _init_state() -> None:
    """Initialise all session-state keys on first run."""
    defaults = {
        "url_history": {},       # url -> {"summary": str, "vector_store": FAISS, "chunks": int}
        "active_urls": [],       # urls currently selected for the merged index
        "merged_store": None,    # FAISS — merged index of all active URLs
        "chat_history": [],      # [{"role": "user"|"assistant", "content": str}]
    }
    for key, val in defaults.items():
        st.session_state.setdefault(key, val)


def _add_url_to_session(url: str, summary: str, vector_store: FAISS, num_chunks: int) -> None:
    """Store a newly processed URL in history and activate it."""
    st.session_state["url_history"][url] = {
        "summary": summary,
        "vector_store": vector_store,
        "chunks": num_chunks,
    }
    if url not in st.session_state["active_urls"]:
        st.session_state["active_urls"].append(url)
    _rebuild_merged_store()


def _rebuild_merged_store() -> None:
    """
    Re-merge the FAISS indexes of all currently active URLs.
    Called whenever the active URL set changes.
    """
    active = st.session_state["active_urls"]
    history = st.session_state["url_history"]
    stores = [history[u]["vector_store"] for u in active if u in history]
    if stores:
        st.session_state["merged_store"] = merge_vector_stores(stores)
    else:
        st.session_state["merged_store"] = None


# ---------------------------------------------------------------------------
# Download helpers
# ---------------------------------------------------------------------------

def _build_pdf(text: str, title: str = "Content Summary") -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_draw_color(37, 99, 235)
    pdf.set_line_width(0.8)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(6)
    pdf.set_font("Helvetica", size=11)
    pdf.set_text_color(30, 41, 59)
    pdf.multi_cell(0, 7, text)
    return bytes(pdf.output())


def _build_docx(text: str, title: str = "Content Summary") -> bytes:
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for para in text.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_download(summary: str) -> None:
    st.markdown("#### ⬇️ Download Summary")
    fmt = st.radio("fmt", ["TXT", "PDF", "DOCX"], horizontal=True, label_visibility="collapsed")
    if fmt == "TXT":
        st.download_button("Download .txt", data=summary,
                           file_name="summary.txt", mime="text/plain",
                           use_container_width=True)
    elif fmt == "PDF":
        try:
            st.download_button("Download .pdf", data=_build_pdf(summary),
                               file_name="summary.pdf", mime="application/pdf",
                               use_container_width=True)
        except Exception as exc:
            st.error(f"PDF error: {exc}")
    elif fmt == "DOCX":
        try:
            st.download_button(
                "Download .docx", data=_build_docx(summary),
                file_name="summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as exc:
            st.error(f"DOCX error: {exc}")


# ---------------------------------------------------------------------------
# UI — Sidebar
# ---------------------------------------------------------------------------

def render_sidebar(groq_api_key_ref: list) -> None:
    """
    Sidebar renders:
      1. API key input
      2. URL input + Add button (supports multiple URLs)
      3. URL history list with checkboxes to activate/deactivate
      4. Active source indicator
    """
    with st.sidebar:
        st.header("⚙️ Configuration")
        st.markdown(
            "Get your free key at [console.groq.com](https://console.groq.com).",
            unsafe_allow_html=True,
        )
        api_key = st.text_input(
            "Groq API Key", type="password",
            placeholder="gsk_...",
            help="Your key is never stored or logged.",
        )
        groq_api_key_ref[0] = api_key

        st.divider()
        st.caption(f"Model: `{DEFAULT_MODEL}`")
        st.caption("Embeddings: `all-MiniLM-L6-v2`")
        st.caption("Supports: YouTube · Web pages")

        # ── Add new URL ──
        st.divider()
        st.markdown("**➕ Add a URL**")
        new_url = st.text_input(
            "URL", placeholder="https://...",
            label_visibility="collapsed",
            key="sidebar_url_input",
        )
        add_clicked = st.button("Load & Index", use_container_width=True, type="primary")

        if add_clicked and new_url.strip():
            if not api_key.strip():
                st.error("Enter your Groq API key first.")
            elif not validators.url(new_url.strip()):
                st.error("Not a valid URL.")
            elif new_url.strip() in st.session_state["url_history"]:
                # Already cached — just activate it
                url = new_url.strip()
                if url not in st.session_state["active_urls"]:
                    st.session_state["active_urls"].append(url)
                    _rebuild_merged_store()
                st.success("URL already indexed — activated!")
            else:
                _process_new_url(new_url.strip(), api_key)

        # ── URL History ──
        if st.session_state["url_history"]:
            st.divider()
            st.markdown("**📚 Loaded URLs**")
            st.caption("Check to include in chat context.")

            for url, data in st.session_state["url_history"].items():
                short = url if len(url) <= 40 else url[:37] + "…"
                is_active = url in st.session_state["active_urls"]
                checked = st.checkbox(
                    f"{short}",
                    value=is_active,
                    key=f"cb_{url}",
                    help=f"{data['chunks']} chunks · {url}",
                )
                if checked and url not in st.session_state["active_urls"]:
                    st.session_state["active_urls"].append(url)
                    _rebuild_merged_store()
                    st.rerun()
                elif not checked and url in st.session_state["active_urls"]:
                    st.session_state["active_urls"].remove(url)
                    _rebuild_merged_store()
                    st.rerun()

            # Clear all
            if st.button("🗑️ Clear all history", use_container_width=True):
                for key in ["url_history", "active_urls", "merged_store", "chat_history"]:
                    st.session_state[key] = {} if key == "url_history" else ([] if key != "merged_store" else None)
                st.rerun()


def _process_new_url(url: str, api_key: str) -> None:
    """Load, summarize and index a new URL. Shows progress in the sidebar."""
    with st.status(f"Processing…", expanded=True) as status:
        try:
            st.write("📥 Loading content…")
            docs = load_documents(url)

            st.write("📄 Splitting into chunks…")
            chunks = split_documents(docs)
            st.write(f"   → {len(chunks)} chunk(s)")

            st.write("🤖 Summarizing…")
            llm = get_llm(api_key)
            summary = run_map_reduce(llm, chunks)

            st.write("🗂️ Building vector index…")
            vector_store = build_vector_store(chunks)

            _add_url_to_session(url, summary, vector_store, len(chunks))
            status.update(label="✅ Done!", state="complete")
            st.rerun()

        except (ValueError, RuntimeError) as exc:
            status.update(label="Failed", state="error")
            st.error(str(exc))
        except Exception as exc:
            status.update(label="Failed", state="error")
            logger.exception("Unexpected error processing %s", url)
            st.error(f"Unexpected error: {exc}")


# ---------------------------------------------------------------------------
# UI — Main panel
# ---------------------------------------------------------------------------

def render_summary_tabs() -> None:
    """
    Show one tab per loaded URL, each with its own summary + download.
    """
    history = st.session_state["url_history"]
    if not history:
        return

    st.markdown("### 📋 Summaries")
    urls = list(history.keys())
    tab_labels = [
        (u if len(u) <= 35 else u[:32] + "…") for u in urls
    ]
    tabs = st.tabs(tab_labels)

    for tab, url in zip(tabs, urls):
        with tab:
            data = history[url]
            st.markdown(
                f"""
                <div style="
                    background:#f0f7ff;
                    border-left:4px solid #2563eb;
                    border-radius:6px;
                    padding:16px 20px;
                    font-size:15px;
                    line-height:1.7;
                    color:#1e293b;
                    margin-bottom:12px;
                ">
                {data["summary"].replace(chr(10), "<br>")}
                </div>
                """,
                unsafe_allow_html=True,
            )
            render_download(data["summary"])


def render_chat_panel(llm: ChatGroq) -> None:
    """
    RAG chat panel. Queries the merged FAISS index (all active URLs).
    Passes the last MEMORY_TURNS turns as conversation history to the LLM.
    """
    merged_store = st.session_state.get("merged_store")
    active_urls  = st.session_state.get("active_urls", [])

    st.markdown("---")
    st.markdown("### 💬 Chat with the Content")

    if not merged_store:
        st.info("Load at least one URL using the sidebar to start chatting.")
        return

    # Active source badge
    source_label = (
        f"{len(active_urls)} URL{'s' if len(active_urls) > 1 else ''} active"
    )
    st.caption(f"🔍 Querying across **{source_label}** · answers are grounded in the loaded content.")

    # Render existing history
    for msg in st.session_state["chat_history"]:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    question = st.chat_input("Ask anything about the loaded content…")
    if not question:
        return

    # Show user message immediately
    st.session_state["chat_history"].append({"role": "user", "content": question})
    with st.chat_message("user"):
        st.markdown(question)

    # Generate and stream answer
    with st.chat_message("assistant"):
        with st.spinner("Thinking…"):
            try:
                answer, sources = answer_question(
                    llm,
                    merged_store,
                    question,
                    st.session_state["chat_history"],
                )
            except Exception as exc:
                answer  = f"❌ Error: {exc}"
                sources = []

        st.markdown(answer)

        if sources:
            with st.expander("📎 Source passages used", expanded=False):
                for i, chunk in enumerate(sources, 1):
                    src = chunk.metadata.get("source", "unknown")
                    short_src = src if len(src) <= 60 else src[:57] + "…"
                    st.markdown(f"**Passage {i}** — `{short_src}`")
                    st.caption(
                        chunk.page_content[:400]
                        + ("…" if len(chunk.page_content) > 400 else "")
                    )

    st.session_state["chat_history"].append({"role": "assistant", "content": answer})


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    st.set_page_config(
        page_title="Summarize & Chat — LangChain + Groq",
        page_icon="🦜",
        layout="wide",
    )

    _init_state()

    # Mutable container so sidebar can write back the API key
    api_key_ref = [""]
    render_sidebar(api_key_ref)
    groq_api_key = api_key_ref[0]

    # ── Header ──
    st.title("🦜 Summarize & Chat with Any URL")
    st.markdown(
        "Add **YouTube links** or **web articles** in the sidebar → get summaries, "
        "download them, and **chat across all sources at once** using RAG with memory."
    )

    has_content = bool(st.session_state["url_history"])

    if not has_content:
        # Empty state — onboarding hint
        st.info(
            "👈 Paste a URL in the sidebar and click **Load & Index** to get started. "
            "You can load multiple URLs and query across all of them.",
            icon="🚀",
        )
        return

    # ── Summaries ──
    render_summary_tabs()

    # ── Chat ──
    if groq_api_key.strip():
        render_chat_panel(get_llm(groq_api_key))
    else:
        st.warning("Enter your Groq API key in the sidebar to enable chat.", icon="🔑")


if __name__ == "__main__":
    main()